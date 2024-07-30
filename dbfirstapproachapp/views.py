import csv
from io import BytesIO
import json
import openpyxl
from docx import Document # type: ignore

import os

from weasyprint import HTML

from django.template.loader import render_to_string
from django.shortcuts import render

from dbfirstapproachapp.models import Categories, Employees, OrderDetails, Orders
import pyodbc
from django.db.models import Q,Avg,Max,Min,Sum,Count
from django.views.decorators.cache import cache_page
from django.core.cache import cache
from django.http import HttpResponse
# Create your views here.

def ShowCategories(request):
    categories = Categories.objects.all();
    return render(request,'dbfirstapproach/showcategories.html',{'categories':categories})

def GetEmployeeBasedOnRawQuery(request):
    sql_query ='''
    SELECT a.OrderID,a.OrderDate,b.CompanyName,c.ProductName,d.UnitPrice,d.Quantity,d.UnitPrice * d.Quantity as 'BillAmount' 
    from Orders a inner join [Order Details] d on a.orderID = d.OrderID 
    inner join Customers b on a.CustomerID=b.CustomerID  inner join  Products c on d.ProductID =c.ProductID where a.OrderID between 10248 and 10255'''
    conn = Connection()
    cursor = conn.cursor()
    cursor.execute(sql_query)
    orders = cursor.fetchall()
    
    return render(request,'dbfirstapproach/showorders.html',{'orders':orders})

#* is efficient way  using procedures
def StoreProceduresDemo(request):
    grandTotal = 0
    runningTotal = 0
    runningOrderTotal = 0
    subtotal = 0
    
    # procedureName ='USP_GetAllOrders'
    conn = Connection()
    cursor = conn.cursor()
    cursor.execute("{call USP_GetAllOrders}")
    orders = cursor.fetchall()
    
    # Calculate runningTotal,runningOrderTotal,grandTotal and Generate Three new Columns
    newOrders =[]
    # related with runningOrderTotal calculate
    previousOrderID =0
    for order in orders:
        
        grandTotal += order.BillAmount
        runningTotal += order.BillAmount

        # calculate runningOrderTotal
        if previousOrderID == 0: 
            previousOrderID = order.OrderID      
            runningOrderTotal += order.BillAmount
            subtotal +=order.BillAmount 
                 
        elif previousOrderID == order.OrderID:
             runningOrderTotal += order.BillAmount
             subtotal +=order.BillAmount
        else:
            newOrders.append(GenerateNewOrder(0,subtotal,0)) # To Calculate subTotal
            previousOrderID = order.OrderID
            runningOrderTotal = order.BillAmount
            subtotal +=order.BillAmount 
            
        newOrders.append(GenerateNewOrder(order,runningTotal,runningOrderTotal))
        
    newOrders.append(GenerateNewOrder(0,subtotal,0))   # subtotal for last order
    return render(request,'dbfirstapproach/showorders.html',{'orders':newOrders,'grandTotal':grandTotal}) 

# Used To Generate order data in addition to two columns (runningTotal,runningOrderTotal)
def GenerateNewOrder (order,runningTotal,runningOrderTotal):
    if (order ==0):
          newOrder = {
            "OrderID":'',
            "OrderDate":'',
            "CompanyName":'',
            "ProductName":'',
            "UnitPrice":'',
            "Quantity":'',
            "BillAmount":'',
            "RunningTotal":runningTotal,
            "RunningOrderTotal":''
        }
    else:
        newOrder = {
            "OrderID":order.OrderID,
            "OrderDate":order.OrderDate,
            "CompanyName":order.CompanyName,
            "ProductName":order.ProductName,
            "UnitPrice":order.UnitPrice,
            "Quantity":order.Quantity,
            "BillAmount":order.BillAmount,
            "RunningTotal":runningTotal,
            "RunningOrderTotal":runningOrderTotal
        }
    return newOrder
  
def SPWithOutputParameter(request):
    # procedureName ='USP_GetAllOrders'
    conn = Connection()
    cursor = conn.cursor()
    # fetch total number of orders with procedures with output parameter
    count=0
    cursor.execute("{call USP_GetOrdersCount(?)}",count)
    counts = cursor.fetchval()
    # # fetch All orders with procedure
    cursor.execute("{call USP_GetAllOrders}") # output parameter => count
    orders = cursor.fetchall()
    
    return render(request,'dbfirstapproach/showorders.html',{'orders':orders,'counts':counts}) 

# Connection to Database by pyodbc
def Connection():
    
    ''' Related With Database Connection'''
    conn = pyodbc.connect('DRIVER=ODBC Driver 17 for SQL Server;Server=KINGDOMTECH;Database=northwind;Trusted_Connection=Yes;')
    return (conn)

def FilteringQuerySetsDemo(request):
    #* Get All Records
    # orders = Orders.objects.all()
    # print(type(orders)) # <class 'django.db.models.query.QuerySet'>
    # print(str(orders.query)) # return raw sql query
    #* Greater Than
    #orders = Orders.objects.filter(freight__gt=20)
    #* Greater Than or equal
    #orders = Orders.objects.filter(freight__gte=20)  
    #* Less Than
    #orders = Orders.objects.filter(freight__lt=20)
    #* Less Than or equal
    #orders = Orders.objects.filter(freight__lte=20)
    #* Exact
    #orders = Orders.objects.filter(shipcountry__exact='Germany')
    #orders = Orders.objects.filter(orderid__exact=10255)
    #* Contains
    #orders = Orders.objects.filter(shipcountry__contains='land')  
    #* in
    #orders = Orders.objects.filter(employeeid__in=[1,2,3,5]).order_by('employeeid') # ascending order
    #orders = Orders.objects.filter(employeeid__in=[1,2,3,5]).order_by('-employeeid') # => - before columnName => descending
    #* startsWith && endsWith
    #orders = Orders.objects.filter(shipname__startswith='A')
    #orders = Orders.objects.filter(shipname__endswith='e') 
    #* Range 
    #orders = Orders.objects.filter(freight__range=[10,20])
    #* two query filter set (OR)
    #orders = Orders.objects.filter(shipname__startswith='A') | Orders.objects.filter(freight__lt=20)
    #* two query filter set using Q model (OR)
    #orders = Orders.objects.filter(Q(shipname__startswith='A') | Q(freight__lt=20))
    #* two query filter set (AND)
    #orders = Orders.objects.filter(shipname__startswith='A') & Orders.objects.filter(freight__lt=20)         
    #* two query filter set using Q model (OR)
    #orders = Orders.objects.filter(Q(shipname__startswith='A') & Q(freight__lt=20))
    #orders = Orders.objects.filter(Q(shipname__startswith='S') & Q(freight__gte=15))
    #* two query filter set (AND) second way
    #orders = Orders.objects.filter(shipname__startswith='A',freight__lt=20)
    
    #! Methods for query sets (Not)
    #* Exclude == ~ ===>  not
    #orders = Orders.objects.exclude(shipname__startswith='A')
    #orders = Orders.objects.filter(~Q(shipname__startswith='S'))
    
    #! Odering 
    #orders = Orders.objects.all().order_by('orderid')
    #orders = Orders.objects.all().order_by('-orderid')
    #orders = Orders.objects.all().order_by('shipcountry')
    #* Filtering Based On Two Columns
    #orders = Orders.objects.all().order_by('employeeid','orderid')
    #* Filter Based on Date =>(year,month,day,hour,second,minute)
    year = 1998
    orders = Orders.objects.filter(orderdate__year=year).order_by('employeeid','orderid')
    
    #! Aggregation
    avg = Orders.objects.all().aggregate(Avg('freight'))
    max = Orders.objects.all().aggregate(Max('freight'))
    min = Orders.objects.all().aggregate(Min('freight'))
    sum = Orders.objects.all().aggregate(Sum('freight'))
    count = Orders.objects.all().aggregate(Count('freight'))
    dict={
        'orders':orders,
        'avg':avg['freight__avg'],
        'max':max['freight__max'],
        'min':min['freight__min'],
        'sum':sum['freight__sum'],
        'count':count['freight__count'],
    }
    
    
    
    return render(request,'dbfirstapproach/filteringDemo.html',dict)
     
def AccordoinDemo(request):
    #* Logic tricky
    orders =Orders.objects.filter(orderid__range=[10248,10255]).order_by('orderid')
    orders_ids = [ order.orderid for order in orders] # ids for all orders
    orders_details = OrderDetails.objects.filter(orderid__in=orders_ids)
    
    dict ={'orders':orders,
           'orders_details':orders_details}
    return render(request,'dbfirstapproach/accordoindemo.html',dict)

def MutiLevelAccordoinDemo(request):
    employees = Employees.objects.filter(employeeid__in=[1,2,4,8])
    employees_ids = [emp.employeeid for emp in employees]
    orders = Orders.objects.filter(employeeid__in=employees_ids,orderdate__month=3,orderdate__day=27);
    orders_ids = [ order.orderid for order in orders] # ids for all orders
    orders_details = OrderDetails.objects.filter(orderid__in=orders_ids)
    dict = {
        'employees':employees,
        'orders':orders,
        'orders_details':orders_details
    }
    return render(request,'dbfirstapproach/mutilevelaccordoin.html',dict)

def ShowOrdersUsingTamplateTag(request):
    return render(request,'dbfirstapproach/showorderusingtemplatetag.html')

# @cache_page(60 * 5) # 5 Minutes Output page cahching (get output of page from cahc memory)
def CachingDemo(request):
    if(cache.get("cache_employees") ==None):
        employees = Employees.objects.all().order_by('employeeid')
        cache.set("cache_employees",employees,3600) # cache for employees
        employees_ids = [emp.employeeid for emp in employees]
        orders = Orders.objects.filter(employeeid__in=employees_ids,orderdate__month=3,orderdate__day=27);
        cache.set("cache_orders",orders,3600) # cache for Orders
        orders_ids = [ order.orderid for order in orders] # ids for all orders
        orders_details = OrderDetails.objects.filter(orderid__in=orders_ids)
        cache.set("cache_orders_details",orders_details,3600) # cache for orders details
    else:
        employees = cache.get('cache_employees')
        orders = cache.get('cache_orders')
        orders_details = cache.get('cache_orders_details')
    
    dict = {
        'employees':employees,
        'orders':orders,
        'orders_details':orders_details
    }
    return render(request,'dbfirstapproach/mutilevelaccordoin.html',dict)

def ExportToCsv(request):
    categories = Categories.objects.all()
    filename = f'Category_data.csv' # filename
    response = HttpResponse(content_type ='text/csv') # create response and defined type of data in response
    '''
    In a regular HTTP response, the Content-Disposition response header is a header indicating if the content is expected to be displayed inline in the browser, that is, as a Web page or as part of a Web page, or as an attachment, that is downloaded and saved locally.
    '''
    response['Content-disposition'] = f'attachment; filename="{filename}"' # key = value =>  filename="{filename}" defined return data 
    writer = csv.writer(response) # used writer object to write data inside response
    
    writer.writerow(['Category Id','Category Name','Description'])
    for category in categories:
        writer.writerow([category.categoryid,
                         category.categoryname,
                         category.description])
    
    return response
                                                                                                     
def ExportToJson(request):
    categories = Categories.objects.all()
    filename = f'Category_data.json' # filename
    response = HttpResponse(content_type ='application/json') # create response and defined type of data in response
    response['Content-disposition'] = f'attachment; filename="{filename}"' # key = value =>  filename="{filename}" defined 
    data = [{'categoryid':category.categoryid,
             'categoryname':category.categoryname,
             'description':category.description} for category in categories] 
    json.dump(data,response) # send data as json format
    return response

# before Export data To Excel Format shloud be install openpyxl library
def ExportToExcel(request):
    categories = Categories.objects.all()
    filename = f'Category_data.xlsx' # filename
    # Generate Excel Workbook 
    workbook = openpyxl.Workbook()
    # Generate Excel Worksheet
    worksheet = workbook.active
    # Add Header
    headers =['Category Id','Category Name','Description']
    worksheet.append(headers)
    # Add Data
    for category in categories:
        worksheet.append([
        category.categoryid,
        category.categoryname,
        category.description
        ])
    # save the workbook to bytesIO buffer
    '''
    StringIO and BytesIO are classes from the io module that provide file-like objects in memory. They act like virtual files, but instead of storing data on an Operating system disk, they store it in memory as strings (StringIO) or bytes (BytesIO).
    '''
    buffer = BytesIO()
    workbook.save(buffer) # save  workbook in buffer
    buffer.seek(0) # 0: means your reference point is the beginning of the file
    
    # Return the Excel File as http response
    response = HttpResponse(buffer.read(),content_type= 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response
    

# before Export data To Word Format shloud be install python-docx library
def ExportToWord(request):
    categories = Categories.objects.all()
    filename = f'Category_data.docx' # filename
    
    # Generate Word File
    document = Document()
    
    # Add a table with headers
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    header_row = table.rows[0].cells
    header_row[0].text = 'Category ID'
    header_row[1].text = 'Category Name'
    header_row[2].text = 'Description'
     
    for cell in header_row:
        cell.paragraphs[0].runs[0].font.bold = True # text bold

    # Add data to the table
    for category in categories:
        row = table.add_row().cells
        row[0].text = str(category.categoryid)
        row[1].text = category.categoryname
        row[2].text = category.description

    # Save the Word document to a BytesIO buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Return the Word file as the HTTP response
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response

# before Export data To Word Format shloud be install weasyprint library and must install GTK and run sever from terminal
def ExportToPDF(request):
    os.add_dll_directory(r"C:\Program Files\GTK3-Runtime Win64\bin")
    categories = Categories.objects.all()
    filename = f'Category_data.pdf' # filename
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    html_string = render_to_string('dbfirstapproach/showcategoriestopdf.html',{'categories':categories})
    html = HTML(string=html_string)
    html.write_pdf(response) # save webpage as string to response in formatting pdf
    
    return response